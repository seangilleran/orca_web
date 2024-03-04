"""TODO: File description."""

import logging
from pathlib import Path

log = logging.getLogger(__name__)


def build_md(images, out_file):
    """TODO: Description."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # We need to load in our base URL in order to make image links.
    # TODO: There's probably a safer way to handle this.
    import os
    from dotenv import load_dotenv

    load_dotenv()

    out_file = Path(out_file)
    out_file.parent.mkdir(parents=True, exist_ok=True)
    out_file_type = out_file.suffix.lower()
    log.info(
        'Writing megadoc (%s) from %d documents to %s...'
        % (out_file_type, len(images), out_file)
    )

    out_file_ic = out_file.with_stem(f"INCOMPLETE_{out_file.stem}")

    album_sizes = {}
    for i, img in enumerate(sorted(images, key=lambda d: d['timestamp'])):
        count_str = '[%d/%d] %s' % (i + 1, len(images), out_file)
        if i == 0 or i % 99 == 0 or i == len(images) - 1:
            log.info(count_str)
        else:
            log.debug(count_str)

        # Find the text file and load its contents.
        txt_file = Path(img['txt_path'])
        if not txt_file.exists():
            log.warning('File not found: %s' % txt_file)
            continue
        with txt_file.open() as f:
            content = f.read()

        # Check album size--we need this to display an iCloud-style heading,
        # and we probably don't want to query the file system every time.
        album_size = album_sizes.get(img['album'], None)
        if not album_size:
            album_size = len([f for f in Path(img['album_path']).glob('*.*')])
            album_sizes[img['album']] = album_size
        album_index = f"{img['album_title']} - {img['index']} of {album_size}"

        # Create a link back to the original image.
        root_url = os.getenv('ORCA_ROOT_URL', '')
        url = f"{root_url}/{img['path']}"

        # Write to Word document--
        if out_file_type == '.docx':
            doc = Document(out_file_ic) if out_file_ic.exists() else Document()

            doc.add_heading(img['timestamp_str'], level=1)

            p = doc.add_paragraph()
            run = p.add_run()
            run.text = f"{album_index}\n"
            run.font.bold = True

            # To add a link we need to manipulate the underlying XML directly.
            run = OxmlElement('w:r')

            link = OxmlElement('w:hyperlink')  # Create link
            link.set(
                qn('r:id'),
                doc.part.relate_to(
                    url,
                    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                    is_external=True,
                ),
            )

            rPr = OxmlElement('w:rPr')  # Format
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            bold = OxmlElement('w:b')
            rPr.append(bold)
            run.append(rPr)

            text_tag = OxmlElement('w:t')  # Set text
            text_tag.text = url
            run.append(text_tag)

            link.append(run)  # Add to paragraph
            p._p.append(link)
            # #####

            doc.add_paragraph('-----')
            doc.add_paragraph(content)

            if i < len(images) - 1:
                doc.add_page_break()

            doc.save(out_file_ic)
            yield i + 1

        # Write to Markdown--
        else:
            with out_file_ic.open('a') as f:
                f.writelines(
                    [
                        '---\n',
                        f"date:  {img['timestamp_str']}\n",
                        f"album: {album_index}\n",
                        f"image: {url}\n",
                        '---\n',
                        '\n',
                        f"{content}\n",
                    ]
                )
                if i < len(images) - 1:
                    f.write('\n\n\n')
            yield i + 1

    # Remove .INCOMPLETE suffix.
    out_file_ic.rename(out_file)


def build_from_search(query_str, batch_path, filetypes=['txt', 'docx']):
    """TODO: Description."""
    import json
    import os
    from orca.search import load_search_cache

    search_index, search_index_file = load_search_cache(batch_path)

    # Get search results and metadata.
    search_info = next(s for s in search_index if s['query_str'] == query_str)
    with Path(search_info['results']['json_path']).open() as f:
        results = json.load(f)

    if not search_info.get('megadocs'):
        search_info['megadocs'] = []
    else:
        # Remove completed filetypes from queue.
        for doc in search_info['megadocs']:
            if doc['complete'] and doc['filetype'] in filetypes:
                filetypes = [f for f in filetypes if f != doc['filetype']]

    if not filetypes or filetypes == []:
        log.info('All megadocs complete or no filetypes specified.')
        return

    # Execute queue.
    doc_path = Path(batch_path) / 'cache' / 'megadocs'
    if not doc_path.exists():
        doc_path.mkdir(exist_ok=True, parents=True)
    stem = Path(search_info['results']['json_path']).stem

    for filetype in filetypes:
        doc_file = doc_path / f"{stem}.{filetype.lower()}"
        megadoc_info = {
            'filetype': filetype,
            'path': f"{doc_file}",
            'pages': 0,
            'size': 0,
            'complete': False,
        }
        search_info['megadocs'].append(megadoc_info)
        
        for i in build_md(results, doc_file):
            megadoc_info['pages'] = i
            with search_index_file.open('w') as f:
                json.dump(search_index, f)
        
        megadoc_info['complete'] = True
        megadoc_info['size'] = os.path.getsize(doc_file)
        with search_index_file.open('w') as f:
                json.dump(search_index, f)


if __name__ == '__main__':
    import argparse

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
    )

    parser = argparse.ArgumentParser()
    parser.add_argument('query')
    parser.add_argument('-b', '--batch_path', required=True)
    args = parser.parse_args()

    results = build_from_search(args.query, args.batch_path)
    log.info('Done!')

